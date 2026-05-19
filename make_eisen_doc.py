from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


doc = Document()

# Base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Title
title = doc.add_heading("Eisen", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

# ---------- 1. Niets wordt geboekt zonder akkoord ----------
doc.add_heading('1.  "Niets wordt geboekt zonder akkoord"', level=1)
p = doc.add_paragraph(style="List Bullet")
p.add_run(
    "Pipeline stopt altijd bij de review-stap. "
)
r = p.add_run("book()")
r.font.name = "Consolas"
p.add_run(
    " wordt alleen aangeroepen als de operator "
)
r = p.add_run("[a]")
r.bold = True
p.add_run(" of ")
r = p.add_run("[fa]")
r.bold = True
p.add_run(" kiest. Escaleren of afsluiten boekt nooit.")

doc.add_paragraph(
    'Bij blocking concerns is [a] uitgeschakeld — voorbeelden van concerns: '
    '"controleer of dit een dubbele factuur is", "Vraag BTW-nummer op".',
    style="List Bullet",
)

# ---------- 2. Liever escaleren dan verkeerd inboeken ----------
doc.add_heading('2.  "Liever escaleren dan verkeerd inboeken"', level=1)
doc.add_paragraph(
    "Bij blocking concern is [a] uitgeschakeld — operator kan alleen "
    "escaleren, corrigeren, of force-approven met expliciete bevestiging.",
    style="List Bullet",
)
doc.add_paragraph(
    "Validator doet aanvullende deterministische checks:",
    style="List Bullet",
)
doc.add_paragraph("ontbrekende verplichte velden → blocking", style="List Bullet 2")
doc.add_paragraph("rekenkundige inconsistentie → warning", style="List Bullet 2")
doc.add_paragraph("toekomstige datum → warning", style="List Bullet 2")

# ---------- 3. Run-log inhoud ----------
doc.add_heading("3.  Run-log per factuur", level=1)
p = doc.add_paragraph()
p.add_run("Per factuur wordt ")
r = p.add_run("runs/<run_id>.json")
r.font.name = "Consolas"
p.add_run(" opgeslagen met:")

rows = [
    ("Wat", "Waar in de log"),
    ("Welk bestand verwerkt", "source_file"),
    ("Voor welke tenant", "tenant_slug"),
    ("Tijdstip", "created_at"),
    ("De exacte system prompt", "steps.extract.system_prompt"),
    (
        "De exacte user prompt (incl. leerregels + voorbeelden die actief waren)",
        "steps.extract.user_prompt",
    ),
    ("De ruwe LLM-respons", "steps.extract.llm_response_raw"),
    ("Confidence per veld", "steps.extract.parsed_fields"),
    ("Agent concerns met reden", "steps.extract.agent_concerns"),
    ("Validator-bevindingen", "steps.validate.concerns"),
    ("Voorgesteld boekstuk", "steps.propose.journal_lines"),
    ("Wat de operator deed", "steps.review.operator_action"),
    ("Welke velden gecorrigeerd", "steps.review.corrections"),
    ("Force-approve bevestiging", "steps.review.operator_confirmation"),
    ("Booking ID", "steps.book_or_escalate.booking_id"),
]

table = doc.add_table(rows=len(rows), cols=2)
table.style = "Light Grid Accent 1"
table.autofit = True

for i, (left, right) in enumerate(rows):
    cell_left = table.rows[i].cells[0]
    cell_right = table.rows[i].cells[1]
    cell_left.text = ""
    cell_right.text = ""

    p_left = cell_left.paragraphs[0]
    run_left = p_left.add_run(left)
    p_right = cell_right.paragraphs[0]
    run_right = p_right.add_run(right)
    run_right.font.name = "Consolas"

    if i == 0:
        run_left.bold = True
        run_right.bold = True
        run_right.font.name = "Calibri"
        shade_cell(cell_left, "D9E1F2")
        shade_cell(cell_right, "D9E1F2")

# ---------- 4. Netwerk & productie-URL ----------
doc.add_heading("4.  Netwerk & productie-URL", level=1)
doc.add_paragraph(
    "In de huidige opzet worden geen echte netwerkverzoeken gedaan.",
    style="List Bullet",
)
p = doc.add_paragraph(style="List Bullet")
p.add_run("De productie-URL (")
r = p.add_run("https://api.boekhouding.nl/v1")
r.font.name = "Consolas"
p.add_run(
    ") is een placeholder; de echte URL komt uit een omgevingsvariabele."
)

# ---------- 5. Operator & authenticatie ----------
doc.add_heading("5.  Operator & authenticatie", level=1)
doc.add_paragraph(
    "Geen authenticatie. De operator is wie het commando uitvoert.",
    style="List Bullet",
)
p = doc.add_paragraph(style="List Bullet")
p.add_run("De CLI is de operator-interface voor MVP. De abstractie in ")
r = p.add_run("review.py")
r.font.name = "Consolas"
p.add_run(
    " maakt een web-UI later mogelijk zonder pipeline-wijzigingen."
)

# ---------- 6. Pipeline ----------
doc.add_heading("6.  Pipeline", level=1)
doc.add_paragraph(
    "Sequentieel, één factuur per run. Geen batch-verwerking of async.",
    style="List Bullet",
)
p = doc.add_paragraph(style="List Bullet")
p.add_run("Run-logs in ")
r = p.add_run("runs/")
r.font.name = "Consolas"
p.add_run(
    " worden niet opgeruimd; bij productiegebruik is een retentiebeleid nodig."
)

out = r"C:\Users\thoma\case\Eisen.docx"
doc.save(out)
print(f"Wrote {out}")
